import os
import docx
from re import compile
from datetime import date
from NoMatchLocated import NoMatchFound
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def main():
    with open(".\\Assignment_Descriptions.txt", 'r') as readfile:
        docs = [string.strip() for string in readfile]

    author = "Eduardo Valles"
    course_professor = "Professor Baker"
    course = "ITEC 4383 Cybersecurity II"
    date = retrieve_date()  # Example date: 'November 23, 2022'
    document_name = ""

    for document_title in docs:
        type = "PT" if "Packet Tracer" in document_title else "Lab"
        assgn_num = locate_asgn_number(document_title) # '26.1.7'

        document_name = "_".join([type, assgn_num, author.split()[0], author.split()[1]]) + '.docx'

        write_document_header(os.path.join(os.path.join(os.getcwd(), "Word_Docs"), document_name),author,
                          course, course_professor, date, document_title.lstrip('-'))

def write_document_header(document_path: str, doc_author: str, course_name: str,
                                    course_professor: str, assignment_date: str, doc_title: str) -> None:
    """
    Accepts a document object and will perform the write operations to standardize each document.
    Here is where you will use the docx module to actually write into the document.

    @param1 The path of the word document to be created and written into
    @param2 The author of the document to be written into
    @param3 The name of the course the document is being created for
    @param4 The name of the professor instructing the course
    @param5 The date the assignment is meant to be due
    @param6 The title of the document that should be centered and bolded
    """
    if not isinstance(doc_author, str): raise TypeError
    if not isinstance(course_professor, str): raise TypeError
    if not isinstance(course_name, str): raise TypeError
    if not isinstance(assignment_date, str): raise TypeError

    document = docx.Document()

    # This portion of the code adds the Header information at the top left of the document
    section = document.sections[0]
    header = section.header

    header_name = header.paragraphs[0]
    header_name.text = f"{doc_author}\n{course_professor}\n{course_name}\n{assignment_date}"

    # This portion of the code writes the title of the assigned centered
    paragraph = document.add_paragraph(doc_title)
    paragraph.add_run().bold = True

    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.save(document_path)

def retrieve_date() -> str:
    """
    A function that returns todays date in Month, Day, Year format.
    Example, 'November 23, 2022'
    """
    d = date.fromordinal(date.toordinal(date.today()))

    formatted_date = d.strftime("%B %d, %Y")

    return formatted_date

def locate_asgn_number(document_title: str) -> str:
    """
    A function designed to locate the Cisco lab numbers.
    """
    assgn_num_regex = compile('\d{1,2}.\d.\d{1,2}')
    mo = assgn_num_regex.search(document_title)

    if mo == None: raise NoMatchFound

    return mo.group()

if __name__ == '__main__':
    main()